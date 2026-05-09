import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_code_block(doc, code_text):
    # 创建一个 1行1列 的无边框表格，用于制作灰色代码底纹
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # 设置浅灰色背景 (F2F2F2)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    tcPr.append(shd)
    
    # 将代码写入单元格，设置字体为 Consolas, 9号 (小五)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(code_text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei') # 中文用微软雅黑
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0) # 黑色字体

def main():
    if not os.path.exists('main.cpp'):
        print("未找到 main.cpp 文件！请确保代码和此脚本在同一文件夹。")
        return

    # 读取 C++ 代码
    with open('main.cpp', 'r', encoding='utf-8') as f:
        code_content = f.read()

    # 创建 Word 文档
    doc = Document()
    
    # 设置全局中文字体支持
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 添加大标题
    title = doc.add_heading('附录', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加子标题
    doc.add_paragraph('支撑材料的代码列表：')
    doc.add_paragraph('(1) 主程序核心算法 main.cpp')
    doc.add_paragraph('(2) 结果可视化脚本 plot_results.py\n')

    # 添加代码块标题
    heading2 = doc.add_heading('程序 1：主程序核心算法 (main.cpp)', level=2)
    
    # 添加排版好的代码框
    print("正在生成排版...")
    add_code_block(doc, code_content)

    # 保存文档
    output_filename = 'Appendix_Code.docx'
    doc.save(output_filename)
    print(f"✅ 生成成功！文件已保存为: {output_filename}")
    print("请在 VSCode 侧边栏下载此 docx 文件发送给写论文的队友。")

if __name__ == '__main__':
    main()